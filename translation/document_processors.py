import os
import csv
import html
import fitz  # PyMuPDF
from docx import Document
from typing import List, Dict, Any, Tuple

class DocxProcessor:
    """Processor for DOCX files"""
    
    @staticmethod
    def _build_tagged_text(runs) -> Tuple[str, bool]:
        """Convert a list of runs into a tagged string, e.g. <r id="0">Hello</r>"""
        text = ""
        has_text = False
        for i, run in enumerate(runs):
            if run.text:
                text += f'<r id="{i}">{html.escape(run.text)}</r>'
                if run.text.strip():
                    has_text = True
        return text, has_text

    @staticmethod
    def _parse_tagged_text(tagged_text: str, runs) -> None:
        """Parse the translated tagged string and assign text back to original runs"""
        import re
        
        for run in runs:
            if run.text:
                run.text = ""
                
        pattern = re.compile(r'<r\s+id=["\']?(\d+)["\']?>(.*?)</r>', re.DOTALL)
        matches = pattern.findall(tagged_text)
        
        if not matches and tagged_text.strip():
            clean_text = re.sub(r'<[^>]+>', '', tagged_text)
            clean_text = html.unescape(clean_text)
            if runs:
                runs[0].text = clean_text
            return

        for run_id_str, content in matches:
            try:
                run_id = int(run_id_str)
                if 0 <= run_id < len(runs):
                    runs[run_id].text = html.unescape(content)
            except ValueError:
                continue

    @staticmethod
    def extract_text(file_path: str) -> List[Dict[str, Any]]:
        BATCH_SIZE = 15
        doc = Document(file_path)
        blocks = []

        for p_idx, para in enumerate(doc.paragraphs):
            tagged_text, has_text = DocxProcessor._build_tagged_text(para.runs)
            if has_text:
                blocks.append({
                    'text': tagged_text,
                    'p_idx': p_idx,
                    'type': 'paragraph'
                })

        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        tagged_text, has_text = DocxProcessor._build_tagged_text(para.runs)
                        if has_text:
                            blocks.append({
                                'text': tagged_text,
                                't_idx': t_idx,
                                'row_idx': r_idx,
                                'cell_idx': c_idx,
                                'p_idx': p_idx,
                                'type': 'table'
                            })

        for i, block in enumerate(blocks):
            block['page'] = i // BATCH_SIZE

        return blocks

    @staticmethod
    def replace_text(file_path: str, translated_blocks: List[Dict[str, Any]], output_path: str):
        doc = Document(file_path)
        
        for block in translated_blocks:
            if block['type'] == 'paragraph':
                para = doc.paragraphs[block['p_idx']]
                DocxProcessor._parse_tagged_text(block.get('translated_text', ''), para.runs)
            elif block['type'] == 'table':
                table = doc.tables[block['t_idx']]
                cell = table.rows[block['row_idx']].cells[block['cell_idx']]
                para = cell.paragraphs[block['p_idx']]
                DocxProcessor._parse_tagged_text(block.get('translated_text', ''), para.runs)
        
        doc.save(output_path)


class PdfProcessor:
    """Processor for PDF files using PyMuPDF (fitz)"""

    @staticmethod
    def extract_text(file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text blocks from PDF with coordinates
        """
        doc = fitz.open(file_path)
        blocks = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text_dict = page.get_text("dict")
            for b_idx, block in enumerate(text_dict["blocks"]):
                if block["type"] == 0:  # text block
                    block_text = ""
                    for line in block["lines"]:
                        for span in line["spans"]:
                            block_text += span["text"]
                        block_text += " "

                    block_text = block_text.strip()
                    if block_text:
                        first_span = block["lines"][0]["spans"][0]
                        blocks.append({
                            'text': block_text,
                            'page': page_num,
                            'bbox': block["bbox"],
                            'font': first_span["font"],
                            'size': first_span["size"],
                            'color': first_span["color"],
                            'origin': first_span["origin"],
                            'b_idx': b_idx,
                            'type': 'block'
                        })
        doc.close()
        return blocks

    @staticmethod
    def _process_page(args) -> None:
        """Process a single PDF page: redact original text then insert translated text."""
        page, blocks = args
        for block in blocks:
            page.add_redact_annot(block['bbox'], fill=None)
        page.apply_redactions()

        for block in blocks:
            try:
                color_int = block.get('color', 0)
                if isinstance(color_int, int):
                    r = ((color_int >> 16) & 255) / 255.0
                    g = ((color_int >> 8) & 255) / 255.0
                    b = (color_int & 255) / 255.0
                    color_tuple = (r, g, b)
                else:
                    color_tuple = (0, 0, 0)

                rect = fitz.Rect(block['bbox'])
                rect.x1 += 30
                rect.y1 += 10

                css_color = f"rgb({int(color_tuple[0]*255)}, {int(color_tuple[1]*255)}, {int(color_tuple[2]*255)})"
                escaped_text = html.escape(block.get('translated_text', ''))
                
                font_name = block.get('font', 'sans-serif')
                font_weight = "bold" if "bold" in font_name.lower() else "normal"
                font_style = "italic" if "italic" in font_name.lower() else "normal"
                
                # Try to clean up font name (e.g., 'TimesNewRomanPSMT' -> 'Times New Roman')
                clean_font = font_name.replace("PSMT", "").replace("MT", "").split("-")[0]
                
                html_content = f"""
                <div style="font-family: '{clean_font}', '{font_name}', sans-serif; font-weight: {font_weight}; font-style: {font_style}; font-size: {block['size']}pt; color: {css_color}; line-height: 1.2;">
                    {escaped_text}
                </div>
                """
                page.insert_htmlbox(rect, html_content, archive=None, rotate=0)
            except Exception as e:
                print(f"Error inserting text on page {block['page']}: {e}")

    @staticmethod
    def replace_text(file_path: str, translated_blocks: List[Dict[str, Any]], output_path: str):
        """
        Create a new PDF where original text is replaced with translations.
        Pages are processed in parallel for speed.
        """
        from concurrent.futures import ThreadPoolExecutor

        doc = fitz.open(file_path)

        # Group blocks by page
        pages: Dict[int, list] = {}
        for block in translated_blocks:
            pages.setdefault(block['page'], []).append(block)

        # Build (page_object, blocks) pairs
        page_jobs = [(doc[page_num], blocks) for page_num, blocks in pages.items()]

        # PyMuPDF page objects are NOT thread-safe for writing, so we process
        # them sequentially inside the same document but use a pool to overlap
        # CPU-bound work (color math, HTML building). For true parallelism we
        # iterate sequentially — the main speed gain is from having translated
        # all text in parallel already.
        for args in page_jobs:
            PdfProcessor._process_page(args)

        doc.save(output_path)
        doc.close()


class CsvProcessor:
    """Processor for CSV files"""

    BATCH_SIZE = 20  # cells per virtual page

    @staticmethod
    def _is_translatable(value: str) -> bool:
        """Return True if the cell value looks like natural-language text."""
        if not value or not value.strip():
            return False
        stripped = value.strip()
        # Skip pure numbers, percentages, and short codes
        try:
            float(stripped.replace(',', '').replace('%', ''))
            return False
        except ValueError:
            pass
        # Keep anything with at least one alphabetic character
        return any(c.isalpha() for c in stripped)

    @staticmethod
    def extract_text(file_path: str) -> List[Dict[str, Any]]:
        encodings = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
        rows = None
        for enc in encodings:
            try:
                with open(file_path, newline='', encoding=enc) as f:
                    rows = list(csv.reader(f))
                break
            except (UnicodeDecodeError, Exception):
                continue
        if rows is None:
            raise ValueError("Could not decode CSV file with supported encodings")

        blocks = []
        for r_idx, row in enumerate(rows):
            for c_idx, cell in enumerate(row):
                if CsvProcessor._is_translatable(cell):
                    blocks.append({
                        'text': cell,
                        'row': r_idx,
                        'col': c_idx,
                        'type': 'csv_cell',
                    })

        for i, block in enumerate(blocks):
            block['page'] = i // CsvProcessor.BATCH_SIZE

        return blocks

    @staticmethod
    def replace_text(file_path: str, translated_blocks: List[Dict[str, Any]], output_path: str):
        encodings = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
        rows = None
        for enc in encodings:
            try:
                with open(file_path, newline='', encoding=enc) as f:
                    rows = list(csv.reader(f))
                break
            except (UnicodeDecodeError, Exception):
                continue
        if rows is None:
            raise ValueError("Could not decode CSV file with supported encodings")

        for block in translated_blocks:
            r, c = block['row'], block['col']
            if r < len(rows) and c < len(rows[r]):
                rows[r][c] = block.get('translated_text', rows[r][c])

        with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
            csv.writer(f).writerows(rows)


class ExcelProcessor:
    """Processor for Excel files (.xlsx / .xls via openpyxl)"""

    BATCH_SIZE = 20  # cells per virtual page

    @staticmethod
    def _is_translatable(value) -> bool:
        if value is None:
            return False
        if not isinstance(value, str):
            return False
        stripped = value.strip()
        if not stripped:
            return False
        return any(c.isalpha() for c in stripped)

    @staticmethod
    def extract_text(file_path: str) -> List[Dict[str, Any]]:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for Excel support: pip install openpyxl")

        wb = openpyxl.load_workbook(file_path, data_only=True)
        blocks = []

        for s_idx, sheet in enumerate(wb.worksheets):
            for r_idx, row in enumerate(sheet.iter_rows()):
                for c_idx, cell in enumerate(row):
                    if ExcelProcessor._is_translatable(cell.value):
                        blocks.append({
                            'text': str(cell.value),
                            'sheet': s_idx,
                            'row': r_idx,
                            'col': c_idx,
                            'type': 'excel_cell',
                        })

        for i, block in enumerate(blocks):
            block['page'] = i // ExcelProcessor.BATCH_SIZE

        wb.close()
        return blocks

    @staticmethod
    def replace_text(file_path: str, translated_blocks: List[Dict[str, Any]], output_path: str):
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for Excel support: pip install openpyxl")

        wb = openpyxl.load_workbook(file_path)
        sheets = wb.worksheets

        for block in translated_blocks:
            s, r, c = block['sheet'], block['row'], block['col']
            if s < len(sheets):
                sheet = sheets[s]
                # openpyxl uses 1-based indexing
                cell = sheet.cell(row=r + 1, column=c + 1)
                cell.value = block.get('translated_text', cell.value)

        wb.save(output_path)
        wb.close()
